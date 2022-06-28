from docx import Document
from docx.shared import Mm, Pt, Inches, Cm
import text_fill
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE

class DocxForm:

    def __init__(self):
        self.row_count_tz = 0
        self.document = Document()
        section = self.document.sections[-1]
        new_width, new_height = section.page_height, section.page_width
        section.page_width = Mm(210)
        section.page_height = Mm(297)

        sections = self.document.sections
        for section in sections:
            section.top_margin = Cm(1.5)
            section.bottom_margin = Cm(1.5)
            section.left_margin = Cm(1.5)
            section.right_margin = Cm(1.5)

        styles = self.document.styles
        style_title = styles.add_style('Title_my', WD_STYLE_TYPE.PARAGRAPH)
        font_title = style_title.font
        font_title.size = Pt(14)
        font_title.name = 'Times New Roman'
        font_title.bold = True

        style_normal = self.document.styles['Normal']
        font = style_normal.font
        font.name = 'Times New Roman'
        font.size = Pt(12)
        font.bold = False


        style_list_number = self.document.styles['List Number']
        font_list_number = style_list_number.font
        font_list_number.name = 'Times New Roman'
        font_list_number.size = Pt(12)
        font_list_number.bold = False

        paragraph = self.document.add_paragraph(text_fill.title, style='Title_my')
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        paragraph_1 = self.document.add_paragraph(
            text_fill.first_par, style='List Number',
        )
        paragraph_1.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

        self.table_1 = self.document.add_table(rows=1, cols=5, style='Table Grid')
        self.table_1.allow_autofit = False
        hdr_cells_1 = self.table_1.rows[0].cells
        widths = [Inches(0.5), Inches(7), Inches(7), Inches(1), Inches(1)]
        for j in range(4):
            hdr_cells_1[j].width = widths[j]

        for i, item in enumerate([text_fill.table_1_1, text_fill.table_1_2, text_fill.table_1_3, text_fill.table_1_4, text_fill.table_1_5]):
            p = hdr_cells_1[i].paragraphs[0]
            p.add_run(item, style = 'Body Text Char').font.size = Pt(12)
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


        self.paragraph_2 = self.document.add_paragraph(
            text_fill.second_par, style='List Number'
        )
        self.paragraph_2.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY


    def common_fill(self, name_ktru, okpd, ktru, measure, quantity):
        row_count = len(self.table_1.rows)
        row_cells = self.table_1.add_row().cells
        row_cells[0].style = 'Normal'
        row_cells[0].text = f'{row_count}.'
        row_cells[1].text = name_ktru
        row_cells[2].text = f'{okpd} / {ktru}'
        row_cells[3].text = measure
        row_cells[4].text = quantity

    def tz_fill(self, name_ktru, nkmi_description, lack_of_description=None, **tz):
        self.row_count_tz+=1
        par = self.document.add_paragraph('', style='Normal')
        par.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        par.add_run('\t', style='Body Text Char')


        if lack_of_description == 'not':
            par.add_run(text_fill.if_tz_empty, style='Body Text Char')
            par.add_run('\n', style='Body Text Char')

        elif lack_of_description == 'not_enough':
            par.add_run(text_fill.if_tz_not_enough, style='Body Text Char')
            par.add_run('\n', style='Body Text Char')

        par.add_run('\n', style='Body Text Char')
        par.add_run(f'2.{self.row_count_tz} ', style='Body Text Char')
        par.add_run('Описание по НКМИ: ', style='Body Text Char')
        par.add_run(nkmi_description, style='Body Text Char')

        self.document.add_paragraph('', style='Normal').alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

        table_2 = self.document.add_table(rows=1, cols=5, style='Table Grid')
        # self.move_table_after(table_2, self.paragraph_2)
        table_2.allow_autofit = False
        hdr_cells_2 = table_2.rows[0].cells
        widths_2 = [Inches(0.3), Inches(6), Inches(4), Inches(4), Inches(2)]
        for j in range(5):
            hdr_cells_2[j].width = widths_2[j]

        for i, item in enumerate([text_fill.table_2_1, text_fill.table_2_2, text_fill.table_2_3, text_fill.table_2_4, text_fill.table_2_5]):
            p = hdr_cells_2[i].paragraphs[0]
            p.add_run(item, style = 'Body Text Char').font.size = Pt(12)
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


        row_count=len(table_2.rows)

        row_cells = table_2.add_row().cells
        row_cells[0].style = 'Normal'
        row_cells[0].text = f'{row_count}.'
        row_cells[1].text = name_ktru


        for i in tz.items():
            row_count = len(table_2.rows)
            row_cells = table_2.add_row().cells
            row_cells[0].text = f'{row_count}.'
            row_cells[1].text = i[0]
            row_cells[3].text = i[1]
            row_cells[4].text = 'КТРУ'



    def doc_save(self):
        paragraph_3 = self.document.add_paragraph(
            text_fill.third_par, style='List Number'
        )
        paragraph_3.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        self.document.save('ТЗ.docx')

    def move_table_after(self, table, paragraph):
        tbl, p = table._tbl, paragraph._p
        p.addnext(tbl)

# d = DocxForm()
